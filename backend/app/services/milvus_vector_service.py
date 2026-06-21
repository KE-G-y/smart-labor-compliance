"""Document parsing and Milvus vector indexing for LangChain retrieval."""
from __future__ import annotations

import csv
import io
import json
import logging
import re
from dataclasses import dataclass
from datetime import datetime
from functools import lru_cache
from html.parser import HTMLParser
from pathlib import Path
from typing import Optional
from uuid import uuid4

from app.schemas.chat import SourceInfo
from app.security import sanitize_text
from app.services.local_model_service import (
    get_local_embeddings,
    get_local_reranker,
    local_embedding_enabled,
    local_embedding_model_path,
    local_reranker_enabled,
    local_reranker_model_path,
)
from app.services.runtime_config import RuntimeConfig


logger = logging.getLogger(__name__)

HYBRID_VECTOR_FIELDS = ["dense", "sparse"]
HYBRID_SEARCH_PARAMS = [
    {"metric_type": "L2", "params": {"ef": 64}},
    {"metric_type": "BM25", "params": {}},
]
DENSE_SEARCH_PARAMS = {"metric_type": "L2", "params": {"ef": 64}}

SUPPORTED_VECTOR_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
    ".xlsx",
}


class VectorStoreUnavailable(RuntimeError):
    """向量库链路不可用时抛出。

    可能原因包括：文件类型不支持、解析依赖没安装、Embedding API 没配置、
    Milvus 连接失败等。
    """


@dataclass(frozen=True)
class VectorIndexResult:
    """一次文档入库后的结果摘要，给接口响应和质量报告使用。"""

    document_id: str
    title: str
    filename: str
    local_file: str
    characters: int
    chunks: int
    collection: str


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        if data and data.strip():
            self.parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self.parts)


def parse_document_text(path: Path) -> str:
    """把不同格式的文件解析成纯文本。

    Milvus 只能存向量和 metadata，不能直接理解 PDF/Word/Excel。
    所以第一步必须先把文件内容抽成文本，再交给 LangChain 切分和 Embedding。
    """
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_VECTOR_EXTENSIONS:
        raise VectorStoreUnavailable(f"暂不支持解析 {suffix or '无扩展名'} 文件")
    if suffix in {".txt", ".md", ".markdown"}:
        return _read_text_file(path)
    if suffix == ".csv":
        return _read_csv_file(path)
    if suffix in {".html", ".htm"}:
        parser = _HTMLTextExtractor()
        parser.feed(_read_text_file(path))
        return parser.text()
    if suffix == ".pdf":
        return _read_pdf_file(path)
    if suffix == ".docx":
        return _read_docx_file(path)
    if suffix == ".xlsx":
        return _read_xlsx_file(path)
    raise VectorStoreUnavailable(f"暂不支持解析 {suffix} 文件")


TEXT_ENCODINGS = ("utf-8-sig", "utf-8", "utf-16", "utf-16-le", "utf-16-be", "gb18030", "gbk", "big5")


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in TEXT_ENCODINGS:
        try:
            text = raw.decode(encoding)
        except UnicodeDecodeError:
            continue
        if "\x00" in text:
            continue
        return text
    return raw.decode("utf-8", errors="ignore")


def _read_csv_file(path: Path) -> str:
    text = _read_text_file(path)
    rows = []
    for row in csv.reader(io.StringIO(text)):
        rows.append(" | ".join(cell.strip() for cell in row if cell and cell.strip()))
    return "\n".join(row for row in rows if row)


def _read_pdf_file(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise VectorStoreUnavailable("PDF 解析依赖 pypdf 未安装") from exc
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        content = page.extract_text() or ""
        if content.strip():
            pages.append(f"[Page {index}]\n{content}")
    return "\n\n".join(pages)


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise VectorStoreUnavailable("Word 解析依赖 python-docx 未安装") from exc
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_xlsx_file(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise VectorStoreUnavailable("Excel 解析依赖 openpyxl 未安装") from exc
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    lines = []
    try:
        for worksheet in workbook.worksheets:
            lines.append(f"[Sheet] {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    lines.append(" | ".join(cells))
    finally:
        workbook.close()
    return "\n".join(lines)


class MilvusVectorService:
    """负责文档向量化入库和相似度检索。

    入库流程：文件 -> 纯文本 -> 分块 chunk -> Embedding 向量 -> Milvus。
    检索流程：问题 -> Embedding 向量 -> Milvus 相似度搜索 -> SourceInfo。
    """

    def __init__(self, runtime_config: RuntimeConfig):
        self.runtime_config = runtime_config

    @property
    def configured(self) -> bool:
        # 缺少任意关键配置时，不尝试连接外部服务，直接让上层走降级提示。
        has_embedding_provider = bool(
            local_embedding_enabled(self.runtime_config)
            or (
                self.runtime_config.langchain_api_key
                and self.runtime_config.langchain_embedding_model
            )
        )
        return bool(
            has_embedding_provider
            and self.runtime_config.milvus_uri
            and self.runtime_config.milvus_collection
        )

    def index_file(
        self,
        *,
        path: Path,
        filename: str,
        local_file: str,
        tenant_id: int,
        tenant_code: str,
        tenant_name: str,
        title: Optional[str] = None,
        source_id: Optional[int] = None,
        document_id: Optional[str] = None,
        extra_metadata: Optional[dict] = None,
    ) -> VectorIndexResult:
        if not self.configured:
            raise VectorStoreUnavailable("请先配置本地 Embedding 或 LangChain API Key，并确认 Milvus 连接")
        # 1. 先把上传文件解析为纯文本，并做脱敏/空白归一化。
        text = sanitize_text(parse_document_text(path)) or ""
        text = _normalize_whitespace(text)
        if not text:
            raise VectorStoreUnavailable("文档解析后没有可入库的文本内容")

        # 2. Markdown frontmatter 和构建脚本传入的 manifest metadata 会合并。
        # 这样 FAQ001、来源编号、地区、风险等级等信息可以随 chunk 一起进入 Milvus。
        raw_metadata = {
            **_parse_markdown_frontmatter(text),
            **(extra_metadata or {}),
        }
        metadata_document_id = sanitize_text(str(raw_metadata.get("document_id") or "")) or ""
        metadata_title = sanitize_text(str(raw_metadata.get("title") or "")) or ""
        document_id = sanitize_text(document_id) or metadata_document_id or uuid4().hex
        document_title = sanitize_text(title) or metadata_title or Path(filename).stem or "未命名文档"
        document_type = _infer_document_type(raw_metadata, filename=filename, local_file=local_file)
        # 3. metadata 是每个 chunk 的“身份证”。检索返回时，前端展示来源、
        # 多租户隔离、FAQ/普通文档区分都依赖这些字段。
        metadata = {
            "tenant_id": int(tenant_id),
            "tenant_code": tenant_code,
            "tenant_name": tenant_name,
            "source_id": int(source_id) if source_id else 0,
            "document_id": document_id,
            "document_type": document_type,
            "title": document_title[:200],
            "filename": filename[:255],
            "local_file": local_file[:500],
            "indexed_at": datetime.utcnow().isoformat(timespec="seconds"),
        }
        metadata.update(_clean_metadata(raw_metadata))
        if document_type == "faq":
            # FAQ 不再进 MySQL 管理，但在向量库里仍要单独标识，
            # 方便检索结果显示 [FAQ]，并和政策原文/制度文档区分。
            metadata["faq_code"] = (sanitize_text(str(raw_metadata.get("faq_code") or document_id)) or document_id)[:80]
            metadata["category"] = (
                sanitize_text(str(raw_metadata.get("category") or metadata.get("category") or "standard_faq"))
                or "standard_faq"
            )[:100]
            metadata["risk_level"] = (
                sanitize_text(str(raw_metadata.get("risk_level") or metadata.get("risk_level") or "medium"))
                or "medium"
            )[:40]

        # 4. 先分块再入库。chunk 太大影响召回精度，太小又容易丢上下文，
        # 所以大小和重叠量由后台配置控制。
        chunks = self._split_text(text, metadata)
        store = self._vector_store()
        try:
            self._add_documents(store, chunks)
        except Exception as exc:
            if self._vector_search_mode() != "hybrid":
                raise
            logger.warning("Milvus hybrid indexing unavailable; retrying with dense indexing. error=%s", exc)
            self._add_documents(self._vector_store(prefer_hybrid=False), chunks)
        return VectorIndexResult(
            document_id=document_id,
            title=document_title,
            filename=filename,
            local_file=local_file,
            characters=len(text),
            chunks=len(chunks),
            collection=self.runtime_config.milvus_collection,
        )

    def _vector_search_mode(self) -> str:
        return str(getattr(self.runtime_config, "vector_search_mode", "hybrid") or "hybrid")

    def _add_documents(self, store, chunks: list) -> None:
        timeout = max(float(getattr(self.runtime_config, "langchain_timeout_seconds", 45)), 60.0)
        batch_size = 16 if self._vector_search_mode() == "hybrid" else 64
        try:
            store.add_documents(chunks, timeout=timeout, batch_size=batch_size)
        except TypeError as exc:
            if "unexpected keyword" not in str(exc):
                raise
            store.add_documents(chunks)

    def index_faq(
        self,
        *,
        faq,
        tenant_id: int,
        tenant_code: str,
        tenant_name: str,
        extra_metadata: Optional[dict] = None,
    ) -> VectorIndexResult:
        # 兼容旧的“FAQ 对象直接向量化”测试/工具路径。
        # 当前正式运营推荐把 FAQ 整理成 Markdown manifest 后用 index_file 入库。
        if not self.configured:
            raise VectorStoreUnavailable("请先配置本地 Embedding 或 LangChain API Key，并确认 Milvus 连接")

        text = faq_to_vector_text(faq)
        if not text:
            raise VectorStoreUnavailable("FAQ 缺少可入库的问题或答案")
        document_id = faq_document_id(faq)
        title = sanitize_text(getattr(faq, "question", "")) or document_id
        metadata = {
            "tenant_id": int(tenant_id),
            "tenant_code": tenant_code,
            "tenant_name": tenant_name,
            "source_id": 0,
            "document_id": document_id,
            "document_type": "faq",
            "faq_id": int(getattr(faq, "id", 0) or 0),
            "faq_code": sanitize_text(getattr(faq, "faq_code", "")) or "",
            "title": title[:200],
            "filename": f"{document_id}.faq.md"[:255],
            "local_file": f"faq://{tenant_code}/{document_id}"[:500],
            "category": sanitize_text(getattr(faq, "category", "")) or "",
            "region": sanitize_text(getattr(faq, "region", "")) or "",
            "risk_level": sanitize_text(getattr(faq, "risk_level", "")) or "medium",
            "language": sanitize_text(getattr(faq, "language", "")) or "zh-CN",
            "indexed_at": datetime.utcnow().isoformat(timespec="seconds"),
        }
        metadata.update(_clean_metadata(extra_metadata or {}))

        chunks = self._split_text(text, metadata)
        store = self._vector_store()
        self.delete_faq_vectors(faq=faq, tenant_id=tenant_id, store=store)
        try:
            self._add_documents(store, chunks)
        except Exception as exc:
            if self._vector_search_mode() != "hybrid":
                raise
            logger.warning("Milvus hybrid FAQ indexing unavailable; retrying with dense indexing. error=%s", exc)
            dense_store = self._vector_store(prefer_hybrid=False)
            self.delete_faq_vectors(faq=faq, tenant_id=tenant_id, store=dense_store)
            self._add_documents(dense_store, chunks)
        return VectorIndexResult(
            document_id=document_id,
            title=title,
            filename=f"{document_id}.faq.md",
            local_file=f"faq://{tenant_code}/{document_id}",
            characters=len(text),
            chunks=len(chunks),
            collection=self.runtime_config.milvus_collection,
        )

    def delete_faq_vectors(self, *, faq, tenant_id: int, store=None) -> bool:
        """按租户和 FAQ id 删除旧向量，避免重复 FAQ 干扰召回。"""
        if not self.configured:
            return False
        faq_id = int(getattr(faq, "id", 0) or 0)
        if not faq_id:
            return False
        vector_store = store or self._vector_store()
        expr = (
            f'metadata["tenant_id"] == {int(tenant_id)} '
            f'and metadata["document_type"] == "faq" '
            f'and metadata["faq_id"] == {faq_id}'
        )
        result = vector_store.delete(expr=expr)
        return bool(result) if result is not None else True

    def similarity_search(self, question: str, *, tenant_id: int, tenant_code: str) -> list[SourceInfo]:
        """按租户检索最相似的知识片段。

        Milvus 负责找“语义最接近”的 chunk；这里再把结果转成前端可展示的来源对象。
        """
        if not self.configured:
            return []
        query = sanitize_text(question) or ""
        if not query:
            return []
        top_k = self.runtime_config.vector_top_k
        expr = f'metadata["tenant_id"] == {int(tenant_id)}'
        documents = self._hybrid_search(query, top_k=top_k, expr=expr)
        if not documents:
            documents = self._dense_search(query, top_k=top_k, expr=expr, tenant_id=tenant_id, tenant_code=tenant_code)
        documents = self._merge_keyword_signal(query, documents, top_k=self._search_k(top_k))
        documents = self._rerank_documents(query, documents, top_k=top_k)
        return [self._source_info_from_document(item) for item in documents[:top_k]]

    def _hybrid_search(self, query: str, *, top_k: int, expr: str) -> list:
        if self.runtime_config.vector_search_mode != "hybrid":
            return []
        try:
            store = self._vector_store(prefer_hybrid=True)
            results = store.similarity_search_with_score(
                query,
                k=self._search_k(top_k),
                param=HYBRID_SEARCH_PARAMS,
                expr=expr,
                ranker_type="weighted",
                ranker_params={"weights": [0.65, 0.35]},
            )
            return [document for document, _score in results]
        except Exception as exc:
            logger.warning("Milvus hybrid search unavailable; falling back to dense search. error=%s", exc)
            return []

    def _dense_search(self, query: str, *, top_k: int, expr: str, tenant_id: int, tenant_code: str) -> list:
        for dense_field in ("dense", None):
            store = self._vector_store(prefer_hybrid=False, dense_field=dense_field)
            try:
                return store.similarity_search(query, k=self._search_k(top_k), param=DENSE_SEARCH_PARAMS, expr=expr)
            except Exception as exc:
                logger.warning("Milvus tenant-scoped dense search failed; retrying with local filter. field=%s error=%s", dense_field or "vector", exc)
                try:
                    documents = store.similarity_search(query, k=self._search_k(top_k), param=DENSE_SEARCH_PARAMS)
                    return [
                        item
                        for item in documents
                        if item.metadata.get("tenant_id") == tenant_id or item.metadata.get("tenant_code") == tenant_code
                    ]
                except Exception as fallback_exc:
                    logger.warning("Milvus dense fallback failed. field=%s error=%s", dense_field or "vector", fallback_exc)
        return []

    def _merge_keyword_signal(self, query: str, documents: list, *, top_k: int) -> list:
        if len(documents) <= 1:
            return documents
        terms = _tokenize_for_keyword_score(query)
        if not terms:
            return documents[:top_k]
        scored = []
        total = len(documents)
        for index, document in enumerate(documents):
            content = f"{document.page_content} {json.dumps(document.metadata or {}, ensure_ascii=False)}".lower()
            keyword_score = sum(1 for term in terms if term in content)
            vector_score = total - index
            scored.append((keyword_score, vector_score, document))
        scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
        return [document for _keyword_score, _vector_score, document in scored[:top_k]]

    def _embeddings(self):
        """创建 Embedding 客户端，用于把文本变成向量。"""
        if local_embedding_enabled(self.runtime_config):
            path = local_embedding_model_path(self.runtime_config)
            try:
                return get_local_embeddings(str(path))
            except Exception as exc:
                raise VectorStoreUnavailable(f"本地 Embedding 不可用：{path}，请确认依赖和模型文件完整：{exc}") from exc

        try:
            from langchain_openai import OpenAIEmbeddings
        except ImportError as exc:
            raise VectorStoreUnavailable("Embedding 依赖 langchain-openai 未安装") from exc

        if not self.runtime_config.langchain_api_key:
            raise VectorStoreUnavailable("未配置 LangChain API Key，且本地 Embedding 未启用或不可用")
        kwargs = {
            "api_key": self.runtime_config.langchain_api_key,
            "model": self.runtime_config.langchain_embedding_model,
            "timeout": self.runtime_config.langchain_timeout_seconds,
        }
        if self.runtime_config.langchain_base_url:
            kwargs["base_url"] = self.runtime_config.langchain_base_url
        return OpenAIEmbeddings(**kwargs)

    def _search_k(self, top_k: int) -> int:
        if local_reranker_enabled(self.runtime_config):
            return max(top_k * 4, top_k)
        return top_k

    def _rerank_documents(self, query: str, documents: list, *, top_k: int) -> list:
        if not local_reranker_enabled(self.runtime_config) or len(documents) <= 1:
            return documents[:top_k]
        path = local_reranker_model_path(self.runtime_config)
        try:
            return get_local_reranker(str(path)).rerank(query, documents, top_k=top_k)
        except Exception as exc:
            logger.warning("Local reranker unavailable; keeping Milvus order. path=%s error=%s", path, exc)
            return documents[:top_k]

    def _vector_store(self, *, prefer_hybrid: Optional[bool] = None, dense_field: Optional[str] = None):
        """创建 LangChain Milvus 向量库对象。

        它封装了连接 Milvus、写入向量和相似度搜索的底层细节。
        """
        try:
            from langchain_milvus import Milvus
            from pymilvus import connections
        except ImportError as exc:
            raise VectorStoreUnavailable("Milvus 依赖 langchain-milvus/pymilvus 未安装") from exc

        connection_args = {"uri": self.runtime_config.milvus_uri}
        if self.runtime_config.milvus_token:
            connection_args["token"] = self.runtime_config.milvus_token

        def ensure_orm_connection(alias: str) -> None:
            if connections.has_connection(alias):
                return
            if self.runtime_config.milvus_token:
                connections.connect(
                    alias=alias,
                    uri=self.runtime_config.milvus_uri,
                    token=self.runtime_config.milvus_token,
                )
            else:
                connections.connect(alias=alias, uri=self.runtime_config.milvus_uri)

        class CompatibleMilvus(Milvus):
            def _init(self, *args, **kwargs):
                ensure_orm_connection(self.alias)
                return super()._init(*args, **kwargs)

        use_hybrid = self._vector_search_mode() == "hybrid" if prefer_hybrid is None else prefer_hybrid
        store_kwargs = {
            "embedding_function": self._embeddings(),
            "collection_name": self.runtime_config.milvus_collection,
            "collection_description": "Smart labor compliance parsed documents",
            "connection_args": connection_args,
            "auto_id": True,
            "metadata_field": "metadata",
            "search_params": DENSE_SEARCH_PARAMS,
        }
        if use_hybrid:
            store_kwargs.update(_hybrid_store_kwargs())
        elif dense_field:
            store_kwargs["vector_field"] = dense_field
        store = CompatibleMilvus(**store_kwargs)
        ensure_orm_connection(store.alias)
        return store

    def _split_text(self, text: str, metadata: dict) -> list:
        """把长文本切成多个 chunk。

        每个 chunk 都会带同一份 metadata，并额外标上 chunk_index，
        方便回答时说明来源来自第几个片段。
        """
        try:
            from langchain_core.documents import Document
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError as exc:
            raise VectorStoreUnavailable("文本切分依赖 langchain-text-splitters 未安装") from exc

        chunk_size = self.runtime_config.vector_chunk_size
        chunk_overlap = min(self.runtime_config.vector_chunk_overlap, max(chunk_size - 1, 0))
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "；", ";", ".", " ", ""],
        )
        chunks = splitter.split_documents([Document(page_content=text, metadata=metadata)])
        for index, chunk in enumerate(chunks):
            chunk.metadata = {**chunk.metadata, "chunk_index": index}
        return chunks

    def _source_info_from_document(self, document) -> SourceInfo:
        """把 Milvus 返回的 Document 转成前端展示用的来源信息。"""
        metadata = document.metadata or {}
        title = metadata.get("title") or metadata.get("filename") or "Milvus 知识片段"
        filename = metadata.get("filename")
        chunk_index = metadata.get("chunk_index")
        document_type = metadata.get("document_type") or "document"
        label = "FAQ" if document_type == "faq" else "文档"
        prefix = f"[{label}] {filename or title}"
        if chunk_index is not None:
            prefix = f"{prefix} #chunk-{chunk_index}"
        snippet = (sanitize_text(document.page_content) or "")[:260]
        return SourceInfo(title=prefix, url=metadata.get("url") or None, snippet=snippet, source_type=document_type)


@lru_cache(maxsize=1)
def _hybrid_store_kwargs() -> dict:
    try:
        from langchain_milvus.function import BM25BuiltInFunction
    except ImportError as exc:
        raise VectorStoreUnavailable("Milvus 混合检索依赖 langchain-milvus BM25BuiltInFunction") from exc
    return {
        "vector_field": HYBRID_VECTOR_FIELDS,
        "builtin_function": BM25BuiltInFunction(
            input_field_names="text",
            output_field_names="sparse",
            analyzer_params={"tokenizer": "icu"},
            enable_match=True,
        ),
        "search_params": HYBRID_SEARCH_PARAMS,
    }


def _tokenize_for_keyword_score(text: str) -> list[str]:
    cleaned = (sanitize_text(text) or "").lower()
    terms = re.findall(r"[a-z0-9_\-]{2,}|[\u4e00-\u9fff]{2,}", cleaned)
    short_terms = re.findall(r"[\u4e00-\u9fff]", cleaned)
    terms.extend("".join(short_terms[index:index + 2]) for index in range(max(len(short_terms) - 1, 0)))
    seen = set()
    result = []
    for term in terms:
        if term not in seen:
            seen.add(term)
            result.append(term)
    return result[:32]


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _parse_markdown_frontmatter(text: str) -> dict[str, str]:
    """读取 Markdown 顶部的简易 frontmatter。

    知识库整理脚本会在文档头部写入 document_id、kb_category、region 等字段，
    这里解析出来后放入 Milvus metadata。
    """
    if not text.startswith("---"):
        return {}
    end = text.find("\n---", 3)
    if end == -1:
        return {}
    metadata: dict[str, str] = {}
    for raw_line in text[3:end].strip().splitlines():
        if ":" not in raw_line:
            continue
        key, raw_value = raw_line.split(":", 1)
        key = key.strip()
        if not key:
            continue
        value = raw_value.strip()
        if value:
            try:
                decoded = json.loads(value)
                value = str(decoded) if decoded is not None else ""
            except (TypeError, ValueError, json.JSONDecodeError):
                value = value.strip("'\"")
        metadata[key] = value
    return metadata


def _infer_document_type(metadata: dict, *, filename: str, local_file: str) -> str:
    """判断当前文件是普通资料还是 FAQ。

    FAQ 会被标成 document_type=faq；普通政策、制度、办事指南则是 document。
    """
    values = {
        "document_type": str(metadata.get("document_type") or "").strip().lower(),
        "kb_category": str(metadata.get("kb_category") or "").strip().lower(),
        "doc_type": str(metadata.get("doc_type") or "").strip().lower(),
        "document_id": str(metadata.get("document_id") or "").strip().upper(),
    }
    normalized_path = f"{filename}/{local_file}".replace("\\", "/").lower()
    if (
        values["document_type"] == "faq"
        or values["kb_category"] == "standard_faq"
        or ("faq" in values["doc_type"] and "问答" in values["doc_type"])
        or "faq标准问答" in values["doc_type"]
        or values["document_id"].startswith("FAQ")
        or "/faqs/" in normalized_path
        or normalized_path.endswith(".faq.md")
    ):
        return "faq"
    return "document"


def faq_document_id(faq) -> str:
    code = sanitize_text(getattr(faq, "faq_code", "")) or ""
    if code:
        return code[:80]
    faq_id = int(getattr(faq, "id", 0) or 0)
    return f"FAQ_DB_{faq_id}" if faq_id else uuid4().hex


def faq_to_vector_text(faq) -> str:
    question = sanitize_text(getattr(faq, "question", "")) or ""
    answer = sanitize_text(getattr(faq, "answer", "")) or ""
    if not question or not answer:
        return ""
    keywords = getattr(faq, "keywords", None)
    aliases = getattr(faq, "aliases", None)
    source_ids = getattr(faq, "source_ids", None)
    lines = [
        "---",
        f'document_id: "{faq_document_id(faq)}"',
        'document_type: "faq"',
        f'faq_id: "{getattr(faq, "id", "") or ""}"',
        f'faq_code: "{sanitize_text(getattr(faq, "faq_code", "")) or ""}"',
        f'category: "{sanitize_text(getattr(faq, "category", "")) or ""}"',
        f'region: "{sanitize_text(getattr(faq, "region", "")) or ""}"',
        f'risk_level: "{sanitize_text(getattr(faq, "risk_level", "")) or "medium"}"',
        f'language: "{sanitize_text(getattr(faq, "language", "")) or "zh-CN"}"',
        "---",
        "",
        f"# FAQ {faq_document_id(faq)} {question}",
        "",
        "## 问题",
        question,
        "",
        "## 标准答案",
        answer,
    ]
    if aliases:
        lines.extend(["", "## 相似问法", "、".join(str(item) for item in aliases if item)])
    if keywords:
        lines.extend(["", "## 关键词", "、".join(str(item) for item in keywords if item)])
    if source_ids:
        lines.extend(["", "## 关联来源 ID", "、".join(str(item) for item in source_ids if item)])
    lines.extend(
        [
            "",
            "## 使用规则",
            "- FAQ 用于提升常见问法召回率。",
            "- 回答时应优先核对关联来源和官方资料；FAQ 与官方来源冲突时，以最新官方来源和人工复核结果为准。",
        ]
    )
    return _normalize_whitespace("\n".join(lines))


def _clean_metadata(metadata: dict) -> dict:
    """清洗 Milvus metadata。

    目标是：只保留可 JSON 序列化、长度可控、不会覆盖核心字段的内容。
    """
    blocked_keys = {
        "tenant_id",
        "tenant_code",
        "tenant_name",
        "source_id",
        "document_id",
        "document_type",
        "faq_id",
        "faq_code",
        "title",
        "filename",
    }
    cleaned = {}
    for key, value in metadata.items():
        if key in blocked_keys or value is None:
            continue
        if isinstance(value, bool):
            cleaned[key] = value
        elif isinstance(value, int):
            cleaned[key] = value
        elif isinstance(value, float):
            cleaned[key] = value
        elif isinstance(value, (list, tuple, set)):
            cleaned[key] = [str(item)[:500] for item in value if item is not None]
        else:
            text = sanitize_text(str(value)) or ""
            if text:
                cleaned[key] = text[:1000]
    return cleaned
